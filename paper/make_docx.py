"""
Sinh file Word mẫu (SNN-Flood-Paper.docx) theo plan của thầy — format conference paper.
    python make_docx.py            (chạy từ thư mục paper/ hoặc repo root)
Nhúng sẵn Table 1 (bảng Word) + 3 figure (pareto_pooled, tsweep, per_region_heatmap).
"""
import csv
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
RES = os.path.join(HERE, "results")
OUT = os.path.join(HERE, "SNN-Flood-Paper.docx")

doc = Document()
# Font mặc định Times New Roman 11pt
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)


def h(text, size=13, space_before=10):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(4)
    return p


def body(text, italic=False, justify=True):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p


def set_col_widths(table, widths):
    """Ép độ rộng từng cột (inch) để tên model không vỡ dòng."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def add_full_table(table_no):
    """Bảng đầy đủ 25 configs đọc từ results/summary.csv, chèn tại chỗ gọi."""
    csv_path = os.path.join(RES, "summary.csv")
    if not os.path.isfile(csv_path):
        body("[TODO: chạy make_figures.py/summarize.py để có results/summary.csv rồi tạo lại docx.]")
        return
    rows = list(csv.DictReader(open(csv_path)))
    rows.sort(key=lambda r: -float(r["flood_IoU"]))
    hdr = ["Model", "Pooled IoU (±std, n)", "Per-chip", "F1", "pwIoU", "Params", "Energy(mJ)", "Spike%"]
    at = doc.add_table(rows=len(rows) + 1, cols=len(hdr))
    at.style = "Light Grid Accent 1"
    for j, x in enumerate(hdr):
        c = at.cell(0, j); c.text = x
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
            for r in p.runs:
                r.bold = True; r.font.size = Pt(8)
    for i, row in enumerate(rows, start=1):
        snn = row["is_snn"].strip().lower() in ("true", "1")
        spk = f"{float(row['spike_rate'])*100:.1f}" if snn else "—"
        vals = [row["model"],
                f"{float(row['flood_IoU']):.3f}±{float(row['flood_IoU_std']):.3f} (n{row['n']})",
                f"{float(row['flood_IoU_chip']):.3f}", f"{float(row['flood_F1']):.3f}",
                f"{float(row['pw_IoU']):.3f}", f"{float(row['params_M']):.2f}",
                f"{float(row['energy_mJ']):.1f}", spk]
        for j, v in enumerate(vals):
            c = at.cell(i, j); c.text = v
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(8)
    set_col_widths(at, [1.6, 1.45, 0.6, 0.5, 0.55, 0.6, 0.72, 0.53])
    cap = doc.add_paragraph()
    cr = cap.add_run(f"Table {table_no}. Full benchmark — all {len(rows)} configurations "
                     "(mean over seeds; n in parentheses). F1 = flood Dice; pwIoU = permanent-water IoU. "
                     "Sorted by pooled flood-IoU.")
    cr.italic = True; cr.font.size = Pt(9); cap.paragraph_format.space_after = Pt(8)


def figure(fname, caption, width=6.2):
    path = os.path.join(FIG, fname)
    if os.path.isfile(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cr = cap.add_run(caption); cr.italic = True; cr.font.size = Pt(9)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(10)


# ---------- Title block ----------
t = doc.add_paragraph()
tr = t.add_run("Energy-Accuracy Trade-offs of Spiking Neural Networks for SAR "
               "Flood Segmentation: A Systematic Benchmark")
tr.bold = True; tr.font.size = Pt(16)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

a = doc.add_paragraph()
ar = a.add_run("Phi Tran¹, Huong Bui², Hoang Dinh²"); ar.font.size = Pt(12)
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff = doc.add_paragraph()
afr = aff.add_run("¹ [Affiliation 1] — [email]    ² [Affiliation 2] — [emails]    [TODO điền cơ quan + email]")
afr.font.size = Pt(10); afr.italic = True
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------- Abstract ----------
h("Abstract", 12)
body("Rapid flood-extent mapping from Sentinel-1 SAR is critical for disaster response, and "
     "on-board or edge deployment demands energy-efficient models. Spiking Neural Networks (SNNs) "
     "promise low-energy inference through sparse, event-driven computation, yet they have never "
     "been benchmarked for flood segmentation on SAR. We present the first systematic energy–accuracy "
     "benchmark of direct-trained SNNs against strong ANN baselines — CNNs, a Transformer, an "
     "ANN-to-SNN conversion, and an INT8-quantized CNN — for 3-class flood-on-land segmentation "
     "(background / permanent water / flood) on Sen1Floods11, where permanent water is separated "
     "from flood using the JRC Global Surface Water layer. Across 20+ configurations trained with "
     "multiple seeds, we find that an INT8 MobileNet-UNet dominates the accuracy–energy Pareto front "
     "(0.464 IoU at 3.1 mJ, versus its FP32 parent's 0.465 at 62.7 mJ) and significantly outperforms the "
     "best direct-trained SNN on both pooled IoU (≈0.47 vs ≈0.35) and a per-chip paired metric, though "
     "the per-chip gap is small (Wilcoxon p = 0.04, Cohen's d = 0.24). "
     "We verify 13–27% spike sparsity supporting the energy estimate, show ANN-to-SNN conversion "
     "fails even at T = 128, analyse per-region geographic sensitivity, and give practical guidance "
     "for choosing SNN vs. quantized CNN.")
kw = doc.add_paragraph()
kwr = kw.add_run("Keywords: "); kwr.bold = True
kw.add_run("Spiking Neural Networks; SAR flood mapping; Sen1Floods11; energy-efficient "
           "segmentation; quantization; neuromorphic computing.")

# ---------- 1 Introduction ----------
h("1. Introduction")
h("1.1 Motivation", 11, 6)
body("Floods are among the most frequent and costly natural disasters worldwide, and rapid, accurate "
     "flood-extent maps are essential for emergency response, damage assessment and relief planning. "
     "Optical satellite imagery is frequently obscured by the very cloud cover that accompanies flood "
     "events, whereas Synthetic Aperture Radar (SAR) — in particular the Sentinel-1 mission [3] — "
     "penetrates clouds and operates day and night, making it the sensor of choice for operational flood "
     "mapping. The Sen1Floods11 benchmark [1] has enabled a wave of deep-learning flood-segmentation "
     "methods on Sentinel-1 data.")
body("However, state-of-the-art segmentation networks (deep CNNs and Transformers) are computationally "
     "and energetically expensive. For time-critical, in-the-field, or on-board (satellite/UAV) "
     "deployment, inference energy becomes a first-class constraint, motivating a careful study of the "
     "energy–accuracy trade-off across model families. Spiking Neural Networks (SNNs) are a promising "
     "candidate: they encode information in sparse binary spikes and replace expensive multiply-"
     "accumulate operations with cheap accumulate-only operations gated by spikes, and map naturally "
     "onto event-driven neuromorphic hardware such as Loihi [8, 9]. Recent advances in surrogate-"
     "gradient direct training [5, 6] and open frameworks [4] have made deep SNNs practical, and SNNs "
     "have begun to be applied to dense prediction tasks such as semantic segmentation [7]. Yet, to our "
     "knowledge, SNNs have never been systematically benchmarked for flood segmentation on SAR, and "
     "prior SNN energy claims are rarely compared against the strongest efficient-ANN competitor, "
     "INT8 quantization. This paper fills that gap with an honest, statistically rigorous benchmark.")
h("1.2 Contributions", 11, 6)
for c in [
    "First direct-trained SNN benchmark for 3-class flood-on-land segmentation on Sen1Floods11, "
    "with 20+ configurations and a full ablation over timesteps T, learning rate and encoding.",
    "Rigorous statistical comparison via Wilcoxon signed-rank on paired per-chip flood-IoU with "
    "bootstrap CIs and effect sizes: the INT8 CNN significantly outperforms the best SNN on both "
    "pooled and per-chip IoU, but the per-chip gap is small (Cohen's d ≈ 0.24) — quantifying exactly "
    "how far SNNs trail on this task.",
    "Verified energy claim: measured spike sparsity 13–27% confirms the SynOps estimate; and "
    "ANN-to-SNN conversion fails to converge even at T = 128, so direct training is essential.",
    "Per-region analysis revealing geographic sensitivity (SNN relatively strong in Sri-Lanka/Ghana, "
    "weak in Mekong/Pakistan).",
    "Practical guidance: a decision procedure for SNN vs. quantized CNN given energy budget, accuracy "
    "floor and hardware target.",
    "Hardware-fair SNN: we introduce Spiking-MobileNet — a depthwise-separable Spiking U-Net (1.52 M "
    "params) — so the SNN uses the same efficient convolutions as MobileNet, isolating the effect of "
    "spiking computation from that of the backbone in the energy comparison.",
]:
    doc.add_paragraph(c, style="List Number")

# ---------- 2 Related Work ----------
h("2. Related Work")
body("2.1 SAR flood mapping. The Sen1Floods11 dataset [1] provides georeferenced Sentinel-1 chips with "
     "hand-labelled water masks across eleven global flood events and has become a standard benchmark; "
     "the Sentinel-1 mission itself is described in [3]. Subsequent work has applied encoder–decoder CNNs "
     "(e.g., U-Net variants [11, 12]) and semi-supervised learning to improve water IoU. Most methods "
     "target binary water/non-water segmentation. We instead address a harder three-class task "
     "(background / permanent water / flood), separating permanent water from flood using the JRC Global "
     "Surface Water layer [2], which is more informative for operational flood response.")
body("2.2 Spiking neural networks for vision. SNNs communicate via sparse binary spikes and are "
     "attractive for energy-efficient, event-driven inference [9]. Two training paradigms dominate: "
     "ANN-to-SNN conversion [20, 21], and direct training with surrogate gradients [5, 6], the latter "
     "enabled at scale by tools such as SpikingJelly [4] and by techniques for training deeper SNNs "
     "[10]. Beyond image classification, Kim et al. [7] showed SNNs can be trained directly for semantic "
     "segmentation. SNN application to remote sensing and SAR is still nascent, and — to our knowledge — "
     "no prior work benchmarks SNNs for flood segmentation on SAR.")
body("2.3 Quantization for edge deployment. Post-training and quantization-aware methods compress CNNs "
     "to INT8 with little accuracy loss while cutting per-operation energy by roughly an order of "
     "magnitude [16, 17, 18]. INT8 is therefore the strongest efficient-ANN competitor to SNNs on "
     "conventional hardware, yet it is seldom included in SNN energy comparisons; we treat it as a "
     "mandatory baseline. Efficient CNN backbones such as MobileNetV2 [13] further reduce compute via "
     "depthwise-separable convolutions, which we also adopt for a hardware-fair spiking counterpart.")

# ---------- 3 Method ----------
h("3. Method")
figure("fig1_pipeline.png",
       "Figure 1. Pipeline overview: Sentinel-1 SAR → preprocessing (dB clip, normalization, JRC "
       "3-class labels) → segmentation model (Spiking U-Net / CNN / Transformer, shared protocol) → "
       "3-class flood map (background / permanent water / flood).")
body("3.1 Problem formulation. We treat flood mapping as three-class semantic segmentation. Given a "
     "two-channel Sentinel-1 chip x ∈ R^{2×H×W} (VV and VH backscatter in dB), a model predicts, for "
     "every pixel, one of three labels: background (0), permanent water (1) or flood (2). Ground-truth "
     "water comes from the Sen1Floods11 hand labels [1]; permanent water is separated from flood using "
     "the JRC Global Surface Water layer [2] (a pixel labelled water is assigned class 1 where JRC marks "
     "permanent water, else class 2). Invalid/NaN pixels are assigned an ignore-index (−1) and excluded "
     "from both the loss and the metrics. All models are optimised and evaluated under one shared "
     "protocol so that differences reflect the architecture, not the training recipe.")
body("3.2 Spiking U-Net. Our core SNN, SNN-Flood, is a four-level U-Net [11] (7.76 M parameters) in "
     "which every convolutional block is Conv → BatchNorm → Leaky Integrate-and-Fire (LIF) neuron. LIF "
     "neurons use the ATan surrogate gradient [6] (τ = 2.0, detach-reset) and run in multi-step mode via "
     "SpikingJelly [4], processing all T timesteps jointly; BatchNorm is applied over the (T, B) axes, "
     "giving a threshold-dependent normalisation similar to tdBN which stabilises deep-SNN training [10]. "
     "SAR inputs use direct encoding (the analog chip is repeated over T timesteps), and the final-layer "
     "membrane potentials are averaged over T and read out as class logits. The architecture mirrors the "
     "vanilla ANN U-Net for a controlled comparison and is illustrated in Figure 2.")
figure("fig2_architecture.png",
       "Figure 2. Spiking U-Net architecture (4-level encoder–decoder). Each block is Conv → tdBN "
       "(over T,B) → LIF (ATan surrogate, τ=2, detach-reset); red dashed lines are skip connections; "
       "the input is repeated over T timesteps and final membrane potentials are averaged to logits.")
body("3.3 Energy modeling. We estimate inference energy analytically at a 45 nm technology node using "
     "the per-operation costs of Horowitz [19]: a 32-bit floating-point multiply-accumulate (MAC) costs "
     "E_MAC ≈ 4.6 pJ, an INT8 MAC ≈ 0.23 pJ, and a spiking accumulate (AC, incurred only when a spike is "
     "present) costs E_AC ≈ 0.9 pJ. For ANNs, energy = FLOPs × E_MAC, with MACs counted consistently by "
     "ptflops. For SNNs, we count synaptic operations SynOps = Σ_conv (MACs × input spike-rate), "
     "accumulated over the T timesteps, and energy = SynOps × E_AC; the spike-rate is measured empirically "
     "(Section 5.3) rather than assumed. This makes SNN energy depend on genuine activation sparsity, not "
     "on nominal model size, and puts SNNs and quantized CNNs on a common analytical footing.")
body("3.4 Statistical testing protocol. Because accuracy differences between strong models are small, "
     "point estimates alone are insufficient. We report two aggregations of flood-IoU: pooled (a single "
     "confusion matrix accumulated over all test pixels) and per-chip (IoU computed per chip, then "
     "averaged). For significance we apply the Wilcoxon signed-rank test to paired per-chip flood-IoU "
     "(models paired chip-by-chip, seeds averaged), and complement p-values with bootstrap 95% "
     "confidence intervals and paired Cohen's d effect sizes for the key model pairs, so that both "
     "statistical significance and practical magnitude are reported.")

# ---------- 4 Experimental Setup ----------
h("4. Experimental Setup")
body("4.1 Dataset. We use the hand-labelled subset of Sen1Floods11 [1]: 446 chips of 512×512 pixels, "
     "each with two Sentinel-1 channels (VV, VH; in dB), spanning eleven geographically diverse flood "
     "events (Bolivia, Ghana, India, Mekong, Nigeria, Pakistan, Paraguay, Somalia, Spain, Sri-Lanka, "
     "USA). Three-class labels are derived as in Section 3.1 using the JRC permanent-water layer [2]. "
     "This geographic diversity enables the per-region analysis of Section 5.4.")
body("4.2 Preprocessing & splits. dB clipped to [−50,0], normalized to [0,1]; NaN → ignore-index. "
     "One NaN-only chip is dropped (445 valid). Region-stratified split 60/20/20 → 271/87/87 chips "
     "(agreed protocol; larger test set than the earlier 70/20/10 for more reliable evaluation), with "
     "spatially-overlapping chips kept in the same split to avoid leakage, fixed seed 42.")
body("4.3 Baselines. Otsu; U-Net (vanilla / SMP-ResNet34 pretrained / U-Net++); DeepLabV3; "
     "MobileNet-UNet (3 LRs) and MobileNet-UNet INT8 (static PTQ); SegFormer-b2; SNN-Flood "
     "(T∈{1..8,10}, LR sweep); Spiking-MobileNet (T∈{2,4}); ANN2SNN (T∈{32,64,128}). "
     "Shared dataset/loss/metric/protocol.")
body("4.4 Training. Optimizer AdamW (lr = 1e-4, weight-decay = 1e-4; LR sweep also uses 5e-5 and 2e-4), "
     "batch size 32 (8 for SNNs due to the ×T memory factor), up to 100 epochs with early stopping "
     "(patience 15) on validation flood-IoU, gradient clipping (max-norm 5.0), 256×256 random crops for "
     "training (full 512×512 for val/test). Loss = weighted CE (class weights [1, 5, 2], ignore-index −1) "
     "+ Dice + Focal (γ = 2). Each learned config uses 3 seeds (T=6 uses 5); Otsu is deterministic.")

# ---------- 5 Results ----------
h("5. Results")
body("We evaluate all models on the held-out test split under the shared protocol of Section 3–4, "
     "reporting accuracy (pooled and per-chip flood-IoU, flood F1, permanent-water IoU), model size, "
     "estimated energy and — for SNNs — measured spike rate. Unless stated otherwise, values are the "
     "mean over three seeds (T = 6 over five).")
h("5.1 Main comparison (Table 1)", 11, 6)
add_full_table(1)
body("On pooled IoU the accuracy leaders are pretrained ANNs (U-Net++ 0.489, SegFormer 0.488); "
     "MobileNet-UNet INT8 reaches 0.464 at only 3.1 mJ (≈ its FP32 parent, 0.465, at 20× lower energy). "
     "SNN-Flood tops out at ≈0.35 pooled IoU, clearly below the quantized CNN. Per-chip significance is "
     "analysed in §5.2.")

h("5.2 Statistical significance", 11, 6)
body("Under the per-chip metric: U-Net-SMP vs MobileNet-INT8 p = 0.0006 (significant, Cohen's d = 0.30); "
     "MobileNet-INT8 vs SNN-T6 p = 0.043 (significant but small, d = 0.24); SNN-T2 vs SNN-T6 p = 0.056 (n.s.). "
     "Thus the INT8 CNN significantly outperforms the best SNN even per-chip, but with a small effect size — "
     "SNNs trail yet stay close on a per-scene basis, while the pooled-IoU gap is large. (An earlier "
     "70/20/10 split with only 45 test chips lacked the power to detect this gap; the 60/20/20 split with "
     "87 test chips does — motivating the larger test set.)")

h("5.3 Energy & spike-rate analysis", 11, 6)
body("The pooled Pareto frontier is Otsu (0 mJ, 0.13) → MobileNet-INT8 (3.1 mJ, 0.464) → SegFormer / "
     "U-Net++ (98–339 mJ, ≈0.49); SNN-Flood is not Pareto-optimal under pooled IoU. Measured spike "
     "sparsity 13–27% confirms energy is driven by genuine sparsity. Direct-trained SNNs "
     "(19–188 mJ) are far more efficient than ANN2SNN (95–441 mJ).")
figure("pareto_pooled.png",
       "Figure 3. Accuracy–energy Pareto front (pooled flood-IoU vs energy, log scale). MobileNet-INT8 "
       "sits at the low-energy corner; the per-chip variant is in paper/figures/pareto_perchip.png.")

h("5.4 Per-region breakdown", 11, 6)
body("Regional difficulty varies widely across the 11 flood events (Fig. 3), and all models share a "
     "similar ordering — indicating difficulty is largely a property of the data, not the model. SNNs "
     "track the same regional pattern as the CNNs, with no region where they uniquely fail.")
figure("per_region_heatmap.png", "Figure 6. Per-region flood-IoU heatmap (model × region).")

h("5.5 Ablations", 11, 6)
body("Timesteps T. Accuracy is essentially flat across T2–T10 (all ≈0.32–0.35; SNN-T2 vs SNN-T6 n.s.), "
     "so larger T buys no accuracy while energy grows proportionally. Training is unstable across seeds: "
     "several configurations show one seed collapsing to ≈0.10 (std up to 0.10–0.15), and which T is "
     "affected varies between runs — pointing to stochastic seed-level collapse rather than a fixed "
     "bimodal dependence on T.")
body("Learning rate. LR = 2e-4 was best for the CNN (MobileNet-UNet 0.462 → 0.465) but degraded the SNN "
     "(SNN-T2 0.340 → 0.316; SNN-T8 0.342 → 0.332): the winning CNN recipe does not transfer to SNNs. "
     "The symmetric LR sweep also establishes a fair comparison (both families tuned).")
body("Quantization. INT8 static post-training quantization preserves accuracy (−0.001 IoU vs FP32) at "
     "≈20× lower per-operation energy, making it the strongest energy-efficient baseline in this study.")
body("Spiking-MobileNet (hardware-fair SNN). [PENDING — results from the next training run] Because the "
     "vanilla Spiking U-Net uses full convolutions while MobileNet uses depthwise-separable ones, we add "
     "a depthwise-separable Spiking-MobileNet (1.52 M params) and will report it in Table 1 and the "
     "Pareto plot, testing whether — on equal, efficient convolutions — the SNN narrows the energy–accuracy "
     "gap to INT8 MobileNet.", italic=True)
figure("ablation_T.png",
       "Figure 7. Ablation over timesteps T: (a) pooled Flood-IoU (mean±std) is flat across T; "
       "(b) mean spike rate; (c) energy grows with T. Increasing T beyond 2 gives no accuracy gain.")

h("5.6 Qualitative results", 11, 6)
body("Figure 5 compares predictions across five geographic regions. U-Net++ and MobileNet delineate the "
     "permanent-water channels (blue) cleanly, whereas Spiking-UNet produces salt-and-pepper predictions "
     "and largely merges water into the flood class — consistent with its low permanent-water IoU "
     "(Table 1). All models still struggle on the hardest scenes.")
figure("qual_comparison.png",
       "Figure 5. Qualitative comparison across regions (columns: SAR VV | Ground truth | U-Net++ | "
       "MobileNet | Spiking-UNet). MobileNet-INT8 is visually identical to MobileNet FP32 (−0.001 IoU).")

# ---------- 6 Discussion ----------
h("6. Discussion")
body("6.1 Why SNN does not dominate. Under a MAC/SynOps energy model, the SNN's sparse-spike "
     "advantage is outweighed by (i) using full convolutions on a vanilla U-Net vs. MobileNet's "
     "depthwise-separable convs, and (ii) INT8 making each MAC ~20× cheaper. The SNN's event-driven "
     "benefit is only partially captured on von-Neumann hardware.")
body("6.2 When to use SNN vs quantized CNN. On standard edge hardware with INT8 kernels, a quantized "
     "MobileNet-UNet is the clear choice (highest accuracy per mJ, best on both metrics). An SNN becomes "
     "attractive only when the target is neuromorphic hardware (e.g., Loihi) where event-driven execution "
     "realises its full energy benefit, when INT8 MAC arrays are unavailable, or with native event-based "
     "sensing; the per-chip gap to INT8 is small (d = 0.24), so the accuracy cost of that choice is modest. "
     "Figure 4 summarises this as a decision guide.")
figure("fig4_decision_tree.png",
       "Figure 4. Decision guide for choosing between a Spiking U-Net and an INT8 MobileNet-UNet given "
       "the deployment target, energy budget, and available INT8 support.")
body("6.3 Limitations. No deployment on real neuromorphic hardware (energy estimated, not on-chip); "
     "single dataset; modest absolute IoU (harder 3-class SAR-only setting).")

# ---------- 7 Conclusion ----------
h("7. Conclusion")
body("We presented the first systematic energy–accuracy benchmark of direct-trained SNNs for SAR "
     "flood segmentation against CNN, Transformer, ANN2SNN and INT8 baselines under one protocol with "
     "multi-seed statistics. A quantized INT8 CNN significantly outperforms the best SNN on both pooled "
     "and per-chip IoU, though the per-chip gap is small (d = 0.24); ANN2SNN fails; quantized CNNs occupy "
     "the pooled Pareto front. SNNs remain a viable low-energy baseline whose advantage awaits neuromorphic "
     "hardware. Future work: neuromorphic deployment, multi-temporal SAR, event-native sensing.")

# ---------- References ----------
h("References")
for r in [
    "§ Dataset & SAR flood mapping",
    "[1] Bonafilia, D., Tellman, B., Anderson, T., & Issenberg, E. (2020). Sen1Floods11: A georeferenced dataset to train and test deep learning flood algorithms for Sentinel-1. CVPRW, 210-211.",
    "[2] Pekel, J. F., Cottam, A., Gorelick, N., & Belward, A. S. (2016). High-resolution mapping of global surface water and its long-term changes. Nature, 540(7633), 418-422.",
    "[3] Torres, R., Snoeij, P., Geudtner, D., et al. (2012). GMES Sentinel-1 mission. Remote Sensing of Environment, 120, 9-24.",
    "[TODO] 3–5 bài SAR flood mapping bằng deep learning gần đây (2022–2025).",
    "§ Spiking neural networks & neuromorphic computing",
    "[4] Fang, W., Chen, Y., Ding, J., et al. (2023). SpikingJelly: An open-source machine learning infrastructure platform for spike-based intelligence. Science Advances, 9(40), eadi1480.",
    "[5] Wu, Y., Deng, L., Li, G., Zhu, J., Xie, Y., & Shi, L. (2019). Direct training for spiking neural networks: Faster, larger, better. AAAI, 33(01), 1311-1318.",
    "[6] Neftci, E. O., Mostafa, H., & Zenke, F. (2019). Surrogate gradient learning in spiking neural networks. IEEE Signal Processing Magazine, 36(6), 51-63.",
    "[7] Kim, Y., Chough, J., & Panda, P. (2022). Beyond classification: Directly training spiking neural networks for semantic segmentation. Neuromorphic Computing and Engineering, 2(4), 044015.",
    "[8] Davies, M., Srinivasa, N., Lin, T. H., et al. (2018). Loihi: A neuromorphic manycore processor with on-chip learning. IEEE Micro, 38(1), 82-99.",
    "[9] Roy, K., Jaiswal, A., & Panda, P. (2019). Towards spike-based machine intelligence with neuromorphic computing. Nature, 575(7784), 607-617.",
    "[10] Zheng, H., Wu, Y., Deng, L., Hu, Y., & Li, G. (2021). Going deeper with directly-trained larger spiking neural networks. AAAI, 35(12), 11062-11070.",
    "[TODO] 3–5 bài SNN gần đây.",
    "§ CNN / Transformer baselines",
    "[11] Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional networks for biomedical image segmentation. MICCAI, 234-241.",
    "[12] Zhou, Z., Rahman Siddiquee, M. M., Tajbakhsh, N., & Liang, J. (2018). UNet++: A nested U-Net architecture for medical image segmentation. DLMIA, 3-11.",
    "[13] Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). MobileNetV2: Inverted residuals and linear bottlenecks. CVPR, 4510-4520.",
    "[14] Chen, L. C., Zhu, Y., Papandreou, G., Schroff, F., & Adam, H. (2018). Encoder-decoder with atrous separable convolution for semantic image segmentation (DeepLabV3+). ECCV, 801-818.",
    "[15] Xie, E., Wang, W., Yu, Z., Anandkumar, A., Alvarez, J. M., & Luo, P. (2021). SegFormer: Simple and efficient design for semantic segmentation with transformers. NeurIPS, 34, 12077-12090.",
    "[TODO] 3–5 bài kiến trúc gần đây.",
    "§ Quantization",
    "[16] Jacob, B., Kligys, S., Chen, B., et al. (2018). Quantization and training of neural networks for efficient integer-arithmetic-only inference. CVPR, 2704-2713.",
    "[17] Krishnamoorthi, R. (2018). Quantizing deep convolutional networks for efficient inference: A whitepaper. arXiv:1806.08342.",
    "[18] Nagel, M., Fournarakis, M., Amjad, R. A., et al. (2021). A white paper on neural network quantization. arXiv:2106.08295.",
    "[TODO] 1–3 bài quantization gần đây.",
    "§ Energy modeling",
    "[19] Horowitz, M. (2014). 1.1 Computing's energy problem (and what we can do about it). ISSCC, 10-14.",
    "[TODO] 1–3 bài energy modeling gần đây.",
    "§ ANN-to-SNN conversion",
    "[20] Rueckauer, B., Lungu, I. A., Hu, Y., Pfeiffer, M., & Liu, S. C. (2017). Conversion of continuous-valued deep networks to efficient event-driven networks for image classification. Frontiers in Neuroscience, 11, 682.",
    "[21] Deng, S., & Gu, S. (2021). Optimal conversion of conventional artificial neural networks to spiking neural networks. ICLR.",
    "[TODO] 1–3 bài ANN2SNN gần đây.",
]:
    if r.startswith("§"):
        p = doc.add_paragraph(); run = p.add_run(r[1:].strip()); run.bold = True; run.italic = True
        run.font.size = Pt(9.5); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
        continue
    p = doc.add_paragraph(r); p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)

# ---------- Appendix A: Reproducibility (bảng đầy đủ đã lên §5.1 Table 2) ----------
doc.add_page_break()
h("Appendix A. Reproducibility")
body("Code, configs, region-stratified splits and fixed seeds are released at "
     "https://github.com/TranNhatPhi/NCKH-thayHuong (commit f984b59). We keep all experiment versions: "
     "v1 = 70/20/10 split (git tag paper-v1, archived in paper/results_v1_70-20-10/); v2 = 60/20/20 split "
     "(current). Hardware: 1× NVIDIA H100 (training); energy is estimated analytically at 45 nm "
     "(not measured on-chip). Seeds: {0,1,2} (T=6 also {3,4}).", italic=True)

doc.save(OUT)
print(f"Đã tạo {OUT}")
